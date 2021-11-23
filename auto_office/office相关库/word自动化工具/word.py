# -*- coding:utf-8 -*-
"""
pandas DataFrame 转换为word表格
@Time: 2021/9/16 14:44
@Author: zhangtaoning
@Email: 
@File: word.py
@Software: PyCharm

- ps: 这里的word模板一定要用特定自定义的模板，否则可能会出现样式找不到的问题
      如果需要用自己的模板，就需要更改类 `Style` 中的配置
"""
import os
import docx
from docx import Document
from docx.shared import Cm
from win32com.client.gencache import EnsureDispatch
from win32com.client import constants


# 各种样式名
class Style:
    body_text = 'Normal'  # 正文文本格式
    table_text = 'No Spacing'  # 表格文本格式
    picture_caption = '图片题注'
    picture = 'No Spacing'
    table_caption = '表格题注'
    table_style = '三线表'  # 表格样式


class Word:
    """
    word文档对象，在这个对象里进行插入图片、文字、表格的操作

    Parameters
    ----------
    template_path : str, default=None
        模板文档路径，指定后word文档的各种样式都和模板文档一致

    save_path : str, default='tmp.docx'
        word文档保存路径

    Attributes
    ----------
    doc : docx.document.Document
        文档对象，主要是对其进行操作
    """
    def __init__(self, template_path=None, save_path='tmp.docx'):
        self.template_path = template_path
        self.save_path = save_path

        self.doc = Document(template_path)  # 文档对象

    def save(self):
        """ 保存word文档 """
        self.doc.save(self.save_path)

    def load(self, reload=False):
        """
        加载文档
        Parameters
        ----------
        reload : bool, default=False
            是否重新加载模板文档，将会丢弃之前添加的内容
        """
        if reload:
            self.doc = Document(self.template_path)
        else:
            self.doc = Document(self.save_path)

    def add_paragraph(self, text, style=Style.body_text):
        """ 增加一段话，默认样式为正文

        Parameters
        ----------
        text : str
            添加的文本内容

        style : str, default=Style.body_text
            文本样式，默认为正文
        """
        p = self.doc.add_paragraph(text, style=style)
        return p

    def add_picture(self, img_path, title=None, max_width=Cm(14.64)):
        """
        添加图片
        Parameters
        ----------
        img_path : str
            图片路径

        title : str, default=None
            图片题注

        max_width : int
            图片最大宽度，使用docx.shared.Cm() 将厘米转为像素点

        """
        pic = self.doc.add_picture(img_path)

        # 自动缩放
        if pic.width > max_width:
            ratio = max_width / pic.width  # 缩放比例
            pic.width = int(pic.width * ratio)
            pic.height = int(pic.height * ratio)

        # 居中刚刚插入的图片
        last_paragraph = self.doc.paragraphs[-1]
        last_paragraph.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER
        last_paragraph.style = Style.picture
        # 插入题注
        if title:
            self.doc.add_paragraph(title, style=Style.picture_caption)

    def add_table(self, df, title=None):
        """

        Parameters
        ----------
        df : DataFrame
            表格数据
        title : str, default=None
            表格题注

        Returns
        -------

        """
        # 设置表格题注
        if title:
            self.doc.add_paragraph(title, style=Style.table_caption)

        table = self.doc.add_table(rows=1, cols=df.shape[1])
        # 设置表头
        for header_cell, header in zip(table.rows[0].cells, df.columns):
            header_cell.text = str(header)
            header_cell.paragraphs[0].style = Style.table_text

        # 给每个单元格都填充上数据
        for _, row in df.iterrows():
            row_cells = table.add_row().cells
            for i, (_, value) in enumerate(row.iteritems()):
                if type(value) is float:
                    row_cells[i].text = "%.4f" % value
                else:
                    row_cells[i].text = str(value)
                row_cells[i].paragraphs[0].style = Style.table_text

        # 设置表格样式
        table.style = Style.table_style

    def add_heading(self, text, level=0):
        self.doc.add_heading(text, level=level)

    def add_caption(self):
        def selection_add_caption(s, style, caption_label):
            s.Find.Style = doc.Styles(style)
            while s.Find.Execute(
                    FindText="",
                    Forward=True,
                    Wrap=1,
                    Format=True,
                    MatchCase=False,
                    MatchWholeWord=True,
                    MatchAllWordForms=False,
                    MatchSoundsLike=False,
                    MatchWildcards=False, ):
                s.MoveLeft()
                s.InsertCaption(Label=caption_label)

        app = EnsureDispatch('Word.Application')
        doc = app.Documents.Open(os.path.abspath(self.save_path))

        selection_add_caption(app.Selection, '表格题注', '表')
        selection_add_caption(app.Selection, '图片题注', '图')

        doc.Close(constants.wdSaveChanges)


# 将dataframe格式的文件转换为word
def df_to_word(df, template_path='tmp', file_path='tmp.docx'):
    document = Document(template_path)
    table = document.add_table(rows=1, cols=df.shape[1])

    # 设置表头
    for header_cell, header in zip(table.rows[0].cells, df.columns):
        header_cell.text = str(header)
        header_cell.paragraphs[0].style = Style.table_text

    # 给每个单元格都填充上数据
    for _, row in df.iterrows():
        row_cells = table.add_row().cells
        for i, (_, value) in enumerate(row.iteritems()):
            if type(value) is float:
                row_cells[i].text = "%.4f" % value
            else:
                row_cells[i].text = str(value)
            row_cells[i].paragraphs[0].style = Style.table_text

    # 设置表格样式
    table.style = Style.table_style
    document.save(file_path)


def test_df_to_word():
    import pandas as pd
    import config
    import os
    df = pd.DataFrame([[1, 2, 3, 4],
                       [2, 3, 4, 5],
                       [3, 4, 5, 6]], columns=['apple', 'pear', 'banana', 'orange'])

    df_to_word(df, template_path=config.word_template, file_path='tmp.docx')
    assert os.path.isfile('tmp.docx')
    os.remove('tmp.docx')


def test_word():
    import config
    word = Word(template_path=config.word_template)


if __name__ == '__main__':
    # test_df_to_word()
    pass
